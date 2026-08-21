from __future__ import annotations

import hashlib
import json
import re
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]

REQUIRED_EVIDENCE = (
    "v11_physics_factorial_arrival_v4_independent_v2",
    "world_model_counterfactual_v144_ranking_confirmation_v1",
    "world_model_counterfactual_v145_shadow_confirmation_parallel_v2",
    "v147b_r1_recovery_synchronized_4h_development_v2",
    "v148_nominal_steady_4h_development_v1",
    "v149s_steady_stress_4h_development_v1",
    "v149d_rush_baseline_4h_development_v1",
    "v150_graph_vs_flat_confirmation_seed17400",
    "v151_paired_vs_absolute_confirmation_seed18400",
    "v151_utility_sensitivity_v1",
)

MODEL_LINEAGE = {
    "v12_charge": {
        "directory": "pi_gwm_multistep_v12_charge_seed{seed}",
        "filename": "physics_graph_world_model_multistep.pt",
        "hashes": {
            42: "452E78A3770C598B6878F617651070461036D42BEF5D1515B73908E7155D4738",
            43: "88C484309CA5AB9780213258708C76E0C7194D0F9D565E0B6620D6C3AB881887",
            44: "1156EAD3CECBF8D01956FADAF0DABC34FBC0EEC701A3DB770385C0E5D61527BC",
        },
    },
    "v13_future_risk": {
        "directory": "pi_gwm_multistep_v13_multiscale_v2_seed{seed}",
        "filename": "physics_graph_world_model_multistep.pt",
        "hashes": {
            42: "5691E96D7F69ADF811AEE689AE87AFE799BA14D8721E3FAA02F9C865AA1E9B28",
            43: "EC1860DDA920637C81FAB2328CE2459F75C460EEA9CFECDC7E08DBF2817DA2DE",
            44: "3E7F46C762592E0A75219900138285C20159DC21DD91FAED75854FC4159C8318",
        },
    },
    "v14_1_counterfactual": {
        "directory": "pi_gwm_counterfactual_v141_seed{seed}",
        "filename": "physics_graph_world_model_counterfactual.pt",
        "hashes": {
            42: "A8D894FCCD94FD63090B4A81E196CBE246F9032DEDEDA932A7FE4968BFE14B3A",
            43: "FFCEE1B230C34EF7680D4871792DDF05D5456A60E3B3C1A721573FF3A3950C56",
            44: "FA50AC594622F84ABC19211F3D70930B148549A37410D9F2D10C2C5C6F69925F",
        },
    },
    "v15_flat_counterfactual_baseline": {
        "directory": "flat_mlp_counterfactual_v150_seed{seed}",
        "filename": "flat_counterfactual_baseline.pt",
        "hashes": {
            42: "75DEEE7C96401B30F63F5AC437EFFEF215A47A7CB00E8296EE1B8BFBED72CFA0",
            43: "4C4FE2FFF630ADE4AA63FBA599603D72539A8982C9E41FCC6155B57A6419147F",
            44: "88E510A444FDD492EF098AB6A51AEC8A09756A501B45264E8B5134FA322C6621",
        },
    },
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def bib_keys(path: Path) -> set[str]:
    text = path.read_text(encoding="utf-8")
    return set(re.findall(r"@[A-Za-z]+\s*\{\s*([^,\s]+)", text))


def citation_keys(path: Path) -> set[str]:
    text = path.read_text(encoding="utf-8")
    return set(re.findall(r"@([A-Za-z0-9_:\-]+)", text))


def check_manuscript(path: Path, references: set[str]) -> dict[str, object]:
    text = path.read_text(encoding="utf-8")
    citations = citation_keys(path)
    abstract_word_count = None
    if "ENGLISH" in path.name:
        abstract_match = re.search(
            r"## Abstract\s+(.*?)\s+\*\*Keywords:\*\*", text, flags=re.DOTALL
        )
        if abstract_match:
            abstract_word_count = len(
                re.findall(r"\b[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*\b", abstract_match.group(1))
            )
    keyword_label = "Keywords" if "ENGLISH" in path.name else "关键词"
    keyword_match = re.search(rf"\*\*{keyword_label}[:：]?\*\*\s*(.+)", text)
    keyword_count = None
    if keyword_match:
        keyword_count = len(
            [item for item in re.split(r"[;；]", keyword_match.group(1)) if item.strip()]
        )
    forbidden = [
        token
        for token in (
            "24-node",
            "24 node",
            "24个节点",
            "high-fidelity AGV digital twin",
            "S-curve",
            "S 型速度曲线",
        )
        if token in text
    ]
    required_phrases = (
        ("20-node", "20 节点"),
        ("120, 360, and 720", "120、360 和 720"),
        ("offline digital-twin", "offline decision twin", "离线数字孪生", "离线决策孪生"),
        ("V12",),
        ("V14.1",),
    )
    missing_concepts = [
        "/".join(options) for options in required_phrases if not any(option in text for option in options)
    ]
    return {
        "citation_count": len(citations),
        "missing_bib_keys": sorted(citations - references),
        "unused_bib_keys": sorted(references - citations),
        "forbidden_topology_terms": forbidden,
        "missing_required_concepts": missing_concepts,
        "abstract_word_count": abstract_word_count,
        "keyword_count": keyword_count,
    }


def check_docx(path: Path) -> dict[str, object]:
    document = Document(path)
    headings = sum(
        paragraph.style.name.startswith("Heading") or paragraph.style.name == "Title"
        for paragraph in document.paragraphs
    )
    return {
        "paragraphs": len(document.paragraphs),
        "headings": headings,
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "first_paragraph": document.paragraphs[0].text if document.paragraphs else "",
        "last_paragraph": document.paragraphs[-1].text if document.paragraphs else "",
    }


def current_word_documents(manuscript_dir: Path) -> list[Path]:
    english = manuscript_dir / "JMS_MULTISTEP_WORLD_MODEL_ENGLISH_CURRENT.docx"
    chinese = manuscript_dir / "JMS_MULTISTEP_WORLD_MODEL_CHINESE_CURRENT.docx"
    chinese_updated = manuscript_dir / "JMS_MULTISTEP_WORLD_MODEL_CHINESE_CURRENT_UPDATED.docx"

    # Prefer the normal filename once it contains the current eight-figure layout.
    # The alternate filename is used only while Word keeps the normal file locked.
    if chinese.is_file() and check_docx(chinese)["inline_shapes"] >= 8:
        selected_chinese = chinese
    else:
        selected_chinese = chinese_updated
    return [english, selected_chinese]


def main() -> int:
    references = bib_keys(ROOT / "manuscript" / "JMS_REFERENCES_VERIFIED.bib")
    manuscript_dir = ROOT / "manuscript" / "current"
    manuscripts = {
        path.name: check_manuscript(path, references)
        for path in sorted(manuscript_dir.glob("JMS_MULTISTEP_WORLD_MODEL_*_CURRENT.md"))
    }
    word_documents = {
        path.name: check_docx(path)
        for path in current_word_documents(manuscript_dir)
        if path.is_file()
    }
    highlights_path = ROOT / "manuscript" / "JMS_HIGHLIGHTS.md"
    highlights = [
        line[2:].strip()
        for line in highlights_path.read_text(encoding="utf-8").splitlines()
        if line.startswith("- ")
    ] if highlights_path.is_file() else []
    highlights_audit = {
        "exists": highlights_path.is_file(),
        "count": len(highlights),
        "character_counts": [len(item) for item in highlights],
        "all_within_85_characters": bool(highlights) and all(len(item) <= 85 for item in highlights),
    }

    evidence = {
        name: {
            "exists": (ROOT / "experiment_results" / name).is_dir(),
            "file_count": len(list((ROOT / "experiment_results" / name).rglob("*")))
            if (ROOT / "experiment_results" / name).is_dir()
            else 0,
        }
        for name in REQUIRED_EVIDENCE
    }

    models = {}
    for stage, specification in MODEL_LINEAGE.items():
        stage_models = {}
        for seed, expected in specification["hashes"].items():
            checkpoint = (
                ROOT
                / "world_model_runs"
                / specification["directory"].format(seed=seed)
                / specification["filename"]
            )
            actual = sha256(checkpoint) if checkpoint.is_file() else None
            stage_models[str(seed)] = {
                "exists": checkpoint.is_file(),
                "expected_sha256": expected,
                "actual_sha256": actual,
                "hash_matches": actual == expected,
            }
        models[stage] = stage_models

    anylogic_source = (
        ROOT
        / "AGV_DT_AnyLogic_Validation"
        / "Manufacturing_AGV_DT_Validation"
        / "Manufacturing_AGV_DT_Validation.alpx"
    )
    anylogic_results = anylogic_source.with_name("anylogic_validation_results.csv")
    figure_dir = ROOT / "paper_outputs" / "submission_figures_current"
    figures = {}
    for number in range(1, 9):
        matches = sorted(figure_dir.glob(f"figure_{number}_*"))
        extensions = sorted({path.suffix.lower() for path in matches})
        figures[str(number)] = {
            "files": [path.name for path in matches],
            "extensions": extensions,
            "has_png": ".png" in extensions,
            "has_vector": bool({".pdf", ".svg"} & set(extensions)),
        }

    report = {
        "root": str(ROOT),
        "reference_count": len(references),
        "manuscripts": manuscripts,
        "word_documents": word_documents,
        "highlights": highlights_audit,
        "formal_evidence": evidence,
        "frozen_models": models,
        "anylogic": {
            "editable_source_exists": anylogic_source.is_file(),
            "results_csv_exists": anylogic_results.is_file(),
        },
        "submission_figures": figures,
    }

    failures = []
    if len(references) < 50:
        failures.append("Reference library contains fewer than 50 entries")
    if len(manuscripts) != 2 or len(word_documents) != 2:
        failures.append("Current bilingual manuscript pair is incomplete")
    for name, item in manuscripts.items():
        if item["missing_bib_keys"] or item["unused_bib_keys"]:
            failures.append(f"Citation coverage mismatch in {name}")
        if item["forbidden_topology_terms"] or item["missing_required_concepts"]:
            failures.append(f"Claim-language mismatch in {name}")
        if "ENGLISH" in name and (
            item["abstract_word_count"] is None or item["abstract_word_count"] > 250
        ):
            failures.append(f"Abstract exceeds 250 words or is missing in {name}")
        if item["keyword_count"] is None or not 1 <= item["keyword_count"] <= 6:
            failures.append(f"Keyword count is outside 1-6 in {name}")
    if not (
        highlights_audit["exists"]
        and 3 <= highlights_audit["count"] <= 5
        and highlights_audit["all_within_85_characters"]
    ):
        failures.append("Highlights do not satisfy the 3-5 bullet and 85-character limits")
    for name, item in word_documents.items():
        if (
            item["paragraphs"] < 80
            or item["headings"] < 15
            or item["tables"] < 3
            or item["inline_shapes"] != 8
        ):
            failures.append(f"DOCX structure is unexpectedly sparse in {name}")
    for name, item in evidence.items():
        if not item["exists"] or item["file_count"] == 0:
            failures.append(f"Missing evidence package: {name}")
    for stage, stage_models in models.items():
        for seed, item in stage_models.items():
            if not item["hash_matches"]:
                failures.append(f"Checkpoint hash mismatch for {stage} seed {seed}")
    if not anylogic_source.is_file() or not anylogic_results.is_file():
        failures.append("AnyLogic source or result table is missing")
    for number, item in figures.items():
        if not item["has_png"] or not item["has_vector"]:
            failures.append(f"Submission figure {number} lacks PNG or vector output")

    report["failures"] = failures
    report["passed"] = not failures
    output = ROOT / "PROJECT_INTEGRITY_AUDIT.json"
    output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["passed"] else 1


if __name__ == "__main__":
    sys.exit(main())
